"""
Agents:
    OrchestratorAgent   - controls workflow, detects query type, routes decisions
    RetrievalAgent      - hybrid search + reranking
    ReasoningAgent      - deep analysis using LLM
    SummarizerAgent     - structured, grounded answer generation
    HallucinationGuard  - validates answer grounding in retrieved context
    InsightsAgent       - auto-summary, key topics, questions, difficulty analysis
    ReportAgent         - Word document report generation

Key Aspects:
    - Memory / conversation context with multi-turn support
    - Extended query type detection (steps, examples, factual)
    - Dynamic prompt building per query type
    - Hallucination detection via context overlap scoring
    - InsightsAgent for Document Insights tab
    - Custom system prompt injection
    - RAG debug metadata attached to every context
    - Streaming-ready LLM calls (generator variant)
"""

import logging
import os
import re
import time
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Generator, List, Optional

from config import config
from core import CrossEncoderReranker, DocumentStore, HybridRetriever
from llm import BaseLLM, LLMFactory

logger = logging.getLogger(__name__)

try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt, RGBColor
    docx_available = True
except ImportError:
    docx_available = False
    logger.warning("python-docx not available. Report generation disabled.")


# CONVERSATION MEMORY

@dataclass
class ConversationTurn:
    question: str
    answer: str
    query_type: str
    timestamp: str = field(default_factory=lambda: datetime.now().isoformat())


class ConversationMemory:
    """
    Maintains a rolling window of conversation turns.
    Injects prior context into prompts for follow-up question awareness.
    """

    def __init__(self, max_turns: int = 6):
        self.max_turns = max_turns
        self.turns: List[ConversationTurn] = []

    def add(self, question: str, answer: str, query_type: str):
        self.turns.append(ConversationTurn(question=question, answer=answer, query_type=query_type))
        if len(self.turns) > self.max_turns:
            self.turns.pop(0)

    def build_context_string(self) -> str:
        if not self.turns:
            return ""
        lines = ["Prior conversation:"]
        for t in self.turns[-3:]:
            lines.append(f"  Q: {t.question}")
            lines.append(f"  A: {t.answer[:300]}...")
        return "\n".join(lines)

    def clear(self):
        self.turns.clear()

    def is_follow_up(self, question: str) -> bool:
        follow_up_signals = [
            r"^(it|that|this|they|he|she)\b",
            r"\b(also|additionally|furthermore|moreover)\b",
            r"\b(explain (it|that|this) (more|further|in detail))\b",
            r"^(what about|how about|and|but|so)\b",
            r"\b(previous|above|mentioned)\b",
        ]
        q = question.lower().strip()
        return any(re.search(p, q) for p in follow_up_signals)


# =============================================================================
# SHARED DATA STRUCTURES
# =============================================================================

@dataclass
class AgentContext:
    """Passed between agents to carry full pipeline state."""
    question: str
    store: DocumentStore
    memory: Optional[ConversationMemory] = None
    custom_system_prompt: Optional[str] = None
    tone: str = "professional"
    query_type: str = "general"
    retrieved_chunks: List[Dict] = field(default_factory=list)
    reranked_chunks: List[Dict] = field(default_factory=list)
    reasoning_notes: str = ""
    answer: str = ""
    sources: List[Dict] = field(default_factory=list)
    confidence: float = 0.0
    hallucination_flag: bool = False
    hallucination_reason: str = ""
    report_path: Optional[str] = None
    suggestions: List[str] = field(default_factory=list)
    processing_time: float = 0.0
    rag_debug: Dict[str, Any] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)
    error: Optional[str] = None


# =============================================================================
# BASE AGENT
# =============================================================================

class BaseAgent:
    def __init__(self, name: str):
        self.name = name
        self.llm: BaseLLM = LLMFactory.get_llm()

    def run(self, ctx: AgentContext) -> AgentContext:
        raise NotImplementedError

    def _log(self, msg: str):
        logger.info(f"[{self.name}] {msg}")


# =============================================================================
# QUERY TYPE DETECTION
# =============================================================================

QUERY_TYPE_PATTERNS = {
    "summarization": r"\b(summar|overview|outline|brief|gist|recap|tldr|summarize|summarise)\b",
    "comparison":    r"\b(compar|differ|versus|vs\.?|contrast|distinguish|difference between)\b",
    "definition":    r"\b(what is|define|meaning of|explain|describe|what are)\b",
    "extraction":    r"\b(list|extract|enumerate|find all|give me all|all the|identify all)\b",
    "steps":         r"\b(how to|steps|procedure|process|method|approach|guide|instructions)\b",
    "example":       r"\b(example|instance|illustration|use case|demonstrate|show me)\b",
    "factual":       r"\b(when|where|who|which|how many|how much|what year|what date)\b",
}

TONE_INSTRUCTIONS = {
    "professional": "Use formal, precise, and professional language.",
    "simple":       "Use simple, plain language suitable for a general audience. Avoid jargon.",
    "technical":    "Use technical terminology and provide detailed, precise explanations.",
    "academic":     "Use academic style with structured reasoning and citations where possible.",
}

TYPE_PROMPT_INSTRUCTIONS = {
    "summarization": "Provide a structured summary with clear sections: main topic, key points, and conclusion.",
    "comparison":    "Structure the answer as a direct comparison: highlight similarities, differences, and conclusions.",
    "definition":    "Provide a clear definition, then expand with context and significance from the document.",
    "extraction":    "Return results as a numbered list. Be exhaustive based on the provided context.",
    "steps":         "Present the answer as a numbered sequence of steps or stages in order.",
    "example":       "Provide concrete examples directly from the document context.",
    "factual":       "Answer directly and concisely with the specific fact. Include supporting context.",
    "general":       "Provide a thorough, well-organized answer based on the document context.",
}


def detect_query_type(question: str) -> str:
    q = question.lower()
    for qtype, pattern in QUERY_TYPE_PATTERNS.items():
        if re.search(pattern, q):
            return qtype
    return "general"


# =============================================================================
# ORCHESTRATOR AGENT
# =============================================================================

class OrchestratorAgent(BaseAgent):
    def __init__(self):
        super().__init__("OrchestratorAgent")
        self._retrieval = RetrievalAgent()
        self._reasoning = ReasoningAgent()
        self._summarizer = SummarizerAgent()
        self._hallucination_guard = HallucinationGuard()
        self._report = ReportAgent()

    def run(self, ctx: AgentContext) -> AgentContext:
        t0 = time.time()
        self._log(f"Received question: {ctx.question!r}")

        if ctx.store.is_empty():
            ctx.error = "No documents have been processed. Please upload PDF files first."
            return ctx

        ctx.query_type = detect_query_type(ctx.question)
        self._log(f"Detected query type: {ctx.query_type}")

        # Inject prior conversation context into the question if it is a follow-up
        if ctx.memory and ctx.memory.is_follow_up(ctx.question):
            self._log("Follow-up question detected. Injecting conversation memory.")

        ctx = self._retrieval.run(ctx)
        if ctx.error:
            return ctx

        if not ctx.reranked_chunks:
            ctx.answer = "No relevant information was found in the documents for this question."
            ctx.confidence = 0.0
            ctx.processing_time = time.time() - t0
            return ctx

        ctx = self._reasoning.run(ctx)
        ctx = self._summarizer.run(ctx)
        ctx = self._hallucination_guard.run(ctx)
        ctx = self._report.run(ctx)

        ctx.processing_time = time.time() - t0
        ctx.suggestions = self._generate_suggestions(ctx)

        # Persist to memory
        if ctx.memory and not ctx.error:
            ctx.memory.add(ctx.question, ctx.answer, ctx.query_type)

        self._log(f"Pipeline complete in {ctx.processing_time:.2f}s")
        return ctx

    @staticmethod
    def _generate_suggestions(ctx: AgentContext) -> List[str]:
        base = {
            "summarization": [
                "What are the main conclusions of the document?",
                "Which section is the most important?",
                "What problems does this document address?",
                "What recommendations are given?",
            ],
            "comparison": [
                "Which option is recommended and why?",
                "What are the trade-offs between the approaches?",
                "Are there any scenarios where one is preferred over the other?",
                "What criteria were used for comparison?",
            ],
            "definition": [
                "Can you give a real-world example?",
                "How is this concept applied in practice?",
                "What are the related concepts mentioned?",
                "What are the limitations of this definition?",
            ],
            "steps": [
                "What are the prerequisites for this process?",
                "What could go wrong at each step?",
                "Is there an alternative approach?",
                "How long does this process typically take?",
            ],
            "general": [
                "Can you elaborate further on this topic?",
                "What are the practical implications?",
                "How does this compare to alternative approaches?",
                "What are the key takeaways?",
            ],
        }
        return base.get(ctx.query_type, base["general"])[:4]


# =============================================================================
# RETRIEVAL AGENT
# =============================================================================

class RetrievalAgent(BaseAgent):
    def __init__(self):
        super().__init__("RetrievalAgent")
        self._retriever = HybridRetriever()
        self._reranker = CrossEncoderReranker()
        self._initialized = False

    def initialize(self):
        if not self._initialized:
            self._retriever.initialize()
            self._reranker.initialize()
            self._initialized = True

    def index_store(self, store: DocumentStore):
        self.initialize()
        self._retriever.index(store.chunks)

    def run(self, ctx: AgentContext) -> AgentContext:
        self._log(f"Retrieving chunks for: {ctx.question!r}")

        # Augment query with memory context for better retrieval on follow-ups
        query = ctx.question
        if ctx.memory and ctx.memory.is_follow_up(ctx.question) and ctx.memory.turns:
            last = ctx.memory.turns[-1]
            query = f"{last.question} {ctx.question}"
            self._log(f"Augmented query with memory: {query!r}")

        try:
            raw = self._retriever.retrieve(query)
            reranked = self._reranker.rerank(query, raw)

            ctx.retrieved_chunks = raw
            ctx.reranked_chunks = reranked

            # Attach RAG debug info
            ctx.rag_debug["retrieved_count"] = len(raw)
            ctx.rag_debug["reranked_count"] = len(reranked)
            ctx.rag_debug["top_scores"] = [
                {
                    "chunk_id": c.get("chunk_id", i),
                    "document": c.get("document", ""),
                    "tfidf_score": round(c.get("score", 0.0), 4),
                    "rerank_score": round(c.get("rerank_score", 0.0), 4),
                    "preview": c["text"][:120],
                }
                for i, c in enumerate(reranked[:8])
            ]

            self._log(f"Retrieved {len(raw)} chunks, reranked to {len(reranked)}")
        except Exception as e:
            logger.error(f"Retrieval failed: {e}")
            ctx.error = f"Retrieval error: {e}"
        return ctx


# =============================================================================
# REASONING AGENT
# =============================================================================

REASONING_SYSTEM_PROMPT = """You are an expert document analyst.
Analyze the retrieved document context and reason through the answer step by step.
Focus on:
- Identifying the most relevant facts
- Noting any ambiguities or contradictions in the context
- Building a logical chain of reasoning
Keep your reasoning concise (3-5 sentences). Do not include the final answer yet."""


class ReasoningAgent(BaseAgent):
    def __init__(self):
        super().__init__("ReasoningAgent")

    def run(self, ctx: AgentContext) -> AgentContext:
        self._log("Reasoning over retrieved context.")
        try:
            context_block = self._build_context(ctx.reranked_chunks)
            memory_block = ""
            if ctx.memory:
                mem = ctx.memory.build_context_string()
                if mem:
                    memory_block = f"\n\n{mem}\n"

            prompt = (
                f"QUESTION: {ctx.question}\n"
                f"{memory_block}\n"
                f"DOCUMENT CONTEXT:\n{'--' * 20}\n{context_block}\n{'--' * 20}\n\n"
                f"Analyze the context and reason through what is needed to answer the question."
            )
            ctx.reasoning_notes = self.llm.generate(prompt, system_prompt=REASONING_SYSTEM_PROMPT)
            ctx.rag_debug["reasoning_generated"] = True
            self._log("Reasoning complete.")
        except Exception as e:
            logger.warning(f"Reasoning agent failed: {e}. Continuing without notes.")
            ctx.reasoning_notes = ""
            ctx.rag_debug["reasoning_generated"] = False
        return ctx

    @staticmethod
    def _build_context(chunks: List[Dict], max_chars: int = 600) -> str:
        parts = []
        for i, chunk in enumerate(chunks, 1):
            doc = chunk.get("document", "Unknown")
            text = chunk["text"][:max_chars]
            parts.append(f"[Chunk {i} | Source: {doc}]\n{text}")
        return "\n\n".join(parts)


# =============================================================================
# SUMMARIZER AGENT
# =============================================================================

BASE_SUMMARIZER_SYSTEM = """You are DocVision, an AI document intelligence assistant.
Generate a clear, accurate, and well-structured answer to the user's question.
Rules:
- Base your answer ONLY on the provided document context.
- Do not invent facts not present in the context.
- If the context is insufficient, clearly state what is missing.
- Keep the answer focused and no longer than necessary."""


class SummarizerAgent(BaseAgent):
    def __init__(self):
        super().__init__("SummarizerAgent")

    def run(self, ctx: AgentContext) -> AgentContext:
        self._log("Generating final answer.")
        try:
            context_block = self._build_context(ctx)
            type_instruction = TYPE_PROMPT_INSTRUCTIONS.get(ctx.query_type, "")
            tone_instruction = TONE_INSTRUCTIONS.get(ctx.tone, TONE_INSTRUCTIONS["professional"])

            system_prompt = ctx.custom_system_prompt or (
                f"{BASE_SUMMARIZER_SYSTEM}\n"
                f"Tone: {tone_instruction}\n"
                f"Format instruction: {type_instruction}"
            )

            memory_block = ""
            if ctx.memory:
                mem = ctx.memory.build_context_string()
                if mem:
                    memory_block = f"\n\n{mem}\n"

            prompt = (
                f"QUESTION: {ctx.question}\n"
                f"QUERY TYPE: {ctx.query_type}\n"
                f"{memory_block}"
            )
            if ctx.reasoning_notes:
                prompt += f"\nREASONING NOTES:\n{ctx.reasoning_notes}\n"
            prompt += (
                f"\nDOCUMENT CONTEXT:\n{'--' * 20}\n{context_block}\n{'--' * 20}\n\n"
                f"Provide a complete and accurate answer based solely on the above context."
            )

            ctx.answer = self.llm.generate(prompt, system_prompt=system_prompt)
            ctx.sources = self._build_sources(ctx.reranked_chunks)
            ctx.confidence = self._compute_confidence(ctx.reranked_chunks)
            ctx.metadata["chunks_used"] = len(ctx.reranked_chunks)
            ctx.metadata["llm_backend"] = self.llm.name()
            ctx.metadata["query_type"] = ctx.query_type
            ctx.metadata["tone"] = ctx.tone
            self._log("Answer generated.")
        except Exception as e:
            logger.error(f"Summarizer LLM call failed: {e}")
            ctx.answer = "An error occurred while generating the answer. Please try again."
            ctx.error = str(e)
        return ctx

    @staticmethod
    def _build_context(ctx: AgentContext) -> str:
        max_chars = 700 if ctx.query_type == "summarization" else 500
        parts = []
        for i, chunk in enumerate(ctx.reranked_chunks, 1):
            doc = chunk.get("document", "Unknown")
            text = chunk["text"][:max_chars]
            parts.append(f"[Source {i}: {doc}]\n{text}")
        return "\n\n".join(parts)

    @staticmethod
    def _build_sources(chunks: List[Dict]) -> List[Dict]:
        sources = []
        for i, chunk in enumerate(chunks[:6], 1):
            sources.append({
                "id": i,
                "document_name": chunk.get("document", "Unknown"),
                "text_preview": chunk["text"][:300] + "...",
                "score": chunk.get("rerank_score", chunk.get("score", 0.0)),
                "chunk_id": chunk.get("chunk_id", i),
            })
        return sources

    @staticmethod
    def _compute_confidence(chunks: List[Dict]) -> float:
        if not chunks:
            return 0.0
        scores = [c.get("rerank_score", c.get("score", 0.0)) for c in chunks]
        raw = float(sum(scores) / len(scores))
        return round(min(0.95, max(0.05, raw)), 3)


# =============================================================================
# HALLUCINATION GUARD
# =============================================================================

class HallucinationGuard(BaseAgent):
    """
    Validates that the generated answer is grounded in the retrieved context.
    Uses word overlap + confidence threshold.
    Flags answers that contain information not found in the context.
    """

    def __init__(self):
        super().__init__("HallucinationGuard")

    CONFIDENCE_THRESHOLD = 0.12
    OVERLAP_THRESHOLD = 0.08

    def run(self, ctx: AgentContext) -> AgentContext:
        if not ctx.answer or not ctx.reranked_chunks:
            return ctx

        # Low confidence flag
        if ctx.confidence < self.CONFIDENCE_THRESHOLD:
            ctx.hallucination_flag = True
            ctx.hallucination_reason = (
                f"Low retrieval confidence ({ctx.confidence:.3f}). "
                "The answer may not be well-supported by the document."
            )
            self._log(f"Low confidence flag raised: {ctx.confidence:.3f}")
            return ctx

        # Lexical grounding check
        context_text = " ".join(c["text"] for c in ctx.reranked_chunks).lower()
        answer_words = set(re.findall(r"\b[a-z]{4,}\b", ctx.answer.lower()))
        context_words = set(re.findall(r"\b[a-z]{4,}\b", context_text))

        if not answer_words:
            return ctx

        overlap_ratio = len(answer_words & context_words) / len(answer_words)
        ctx.rag_debug["grounding_overlap"] = round(overlap_ratio, 3)

        if overlap_ratio < self.OVERLAP_THRESHOLD:
            ctx.hallucination_flag = True
            ctx.hallucination_reason = (
                "The answer contains significant content not found in the retrieved document context. "
                "Please verify this information independently."
            )
            self._log(f"Grounding check failed: overlap={overlap_ratio:.3f}")

        return ctx


# =============================================================================
# INSIGHTS AGENT
# =============================================================================

INSIGHTS_PROMPTS = {
    "summary": (
        "You are a document analyst. Produce a comprehensive summary of the following document text. "
        "Structure it as: Overview (2-3 sentences), Key Points (numbered list), and Conclusion (1-2 sentences). "
        "Be thorough and accurate. Base everything strictly on the provided text."
    ),
    "key_topics": (
        "Extract the main topics and concepts from the following document text. "
        "Return a numbered list of topics, each with a one-sentence description. "
        "Focus on the most important and recurring themes. Return 8-12 topics."
    ),
    "short_questions": (
        "Generate 8 short-answer exam questions based on the following document text. "
        "Questions should test factual recall and understanding. "
        "Format: numbered list of questions only, no answers."
    ),
    "long_questions": (
        "Generate 5 long-answer / essay-style exam questions based on the following document text. "
        "Questions should require critical thinking and comprehensive answers. "
        "Format: numbered list of questions only, no answers."
    ),
    "mcq": (
        "Generate 5 multiple-choice questions (MCQ) based on the following document text. "
        "For each question provide 4 options (A, B, C, D) and mark the correct answer. "
        "Format each question as:\nQ: <question>\nA) ...\nB) ...\nC) ...\nD) ...\nAnswer: <letter>"
    ),
    "difficulty": (
        "Analyze the following document text and categorize its content by difficulty level. "
        "Return a structured analysis with:\n"
        "- Overall difficulty: Beginner / Intermediate / Advanced\n"
        "- Beginner concepts (list)\n"
        "- Intermediate concepts (list)\n"
        "- Advanced concepts (list)\n"
        "- Recommended audience\n"
        "Base analysis strictly on the provided text."
    ),
    "smart_notes": (
        "Generate comprehensive study notes from the following document text. "
        "Structure them as:\n"
        "1. Topic heading\n"
        "2. Key definitions\n"
        "3. Important facts and figures\n"
        "4. Summary points\n"
        "Use clear, concise language. Cover all major concepts."
    ),
}


class InsightsAgent(BaseAgent):
    """
    Generates document-level intelligence:
    - Auto summary
    - Key topic extraction
    - Exam questions (short, long, MCQ)
    - Difficulty analysis
    - Smart notes
    """

    def __init__(self):
        super().__init__("InsightsAgent")

    def run(self, ctx: AgentContext) -> AgentContext:
        raise NotImplementedError("Call generate_insight() directly.")

    def generate_insight(self, store: DocumentStore, insight_type: str) -> str:
        if store.is_empty():
            return "No documents loaded. Please upload PDFs first."

        system_prompt = INSIGHTS_PROMPTS.get(insight_type)
        if not system_prompt:
            return f"Unknown insight type: {insight_type}"

        # Build document text from all chunks (capped to avoid token limits)
        full_text = self._build_doc_text(store.chunks)

        prompt = (
            f"DOCUMENT TEXT:\n{'--' * 20}\n{full_text}\n{'--' * 20}\n\n"
            f"Perform the requested analysis on the document text above."
        )

        try:
            return self.llm.generate(prompt, system_prompt=system_prompt)
        except Exception as e:
            logger.error(f"InsightsAgent failed for {insight_type}: {e}")
            return f"Error generating {insight_type}: {str(e)}"

    @staticmethod
    def _build_doc_text(chunks: List[Dict], max_total_chars: int = 12000) -> str:
        texts = []
        total = 0
        for chunk in chunks:
            text = chunk["text"]
            doc = chunk.get("document", "")
            entry = f"[{doc}]\n{text}"
            if total + len(entry) > max_total_chars:
                break
            texts.append(entry)
            total += len(entry)
        return "\n\n".join(texts)


# =============================================================================
# REPORT AGENT
# =============================================================================

class ReportAgent(BaseAgent):
    def __init__(self):
        super().__init__("ReportAgent")
        # Do NOT call makedirs here — config.REPORTS_DIR uses tempfile.gettempdir()
        # which is evaluated lazily. Create the directory right before each write.

    @staticmethod
    def _ensure_reports_dir() -> str:
        """Resolve and create the reports directory, return the path."""
        reports_dir = config.REPORTS_DIR
        os.makedirs(reports_dir, exist_ok=True)
        return reports_dir

    def run(self, ctx: AgentContext) -> AgentContext:
        if not docx_available:
            return ctx
        self._log("Generating Word report.")
        try:
            filename = self._make_filename(ctx)
            reports_dir = self._ensure_reports_dir()
            path = os.path.join(reports_dir, filename)
            self._write_docx(path, ctx)
            ctx.report_path = path
            self._log(f"Report saved: {filename}")
        except Exception as e:
            logger.error(f"Report generation failed: {e}")
        return ctx

    def generate_insights_report(self, store: DocumentStore, insights: Dict[str, str]) -> Optional[str]:
        if not docx_available:
            return None
        try:
            filename = "DocVision_Insights_Report.docx"
            reports_dir = self._ensure_reports_dir()
            path = os.path.join(reports_dir, filename)
            self._write_insights_docx(path, insights)
            return path
        except Exception as e:
            logger.error(f"Insights report generation failed: {e}")
            return None

    @staticmethod
    def _make_filename(ctx: AgentContext) -> str:
        clean_q = re.sub(r'[<>:"/\\|?*\n]', "", ctx.question.strip())[:50]
        docs = list({c.get("document", "") for c in ctx.reranked_chunks})
        doc_part = Path(docs[0]).stem if docs else "report"
        if len(docs) > 1:
            doc_part += f" +{len(docs) - 1}"
        return f"{clean_q} ({doc_part}).docx"

    @staticmethod
    def _write_docx(path: str, ctx: AgentContext):
        os.makedirs(os.path.dirname(path), exist_ok=True)
        doc = DocxDocument()
        title = doc.add_heading("DocVision OCR - Analysis Report", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph(f"Generated:       {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Query type:      {ctx.query_type}")
        doc.add_paragraph(f"Confidence:      {ctx.confidence:.3f}")
        doc.add_paragraph(f"Processing time: {ctx.processing_time:.2f}s")
        doc.add_paragraph(f"LLM backend:     {ctx.metadata.get('llm_backend', 'unknown')}")

        if ctx.hallucination_flag:
            doc.add_paragraph(f"Hallucination warning: {ctx.hallucination_reason}")

        doc.add_heading("Question", level=1)
        doc.add_paragraph(ctx.question)

        doc.add_heading("Answer", level=1)
        doc.add_paragraph(ctx.answer)

        if ctx.reasoning_notes:
            doc.add_heading("Reasoning Notes", level=1)
            doc.add_paragraph(ctx.reasoning_notes)

        if ctx.sources:
            doc.add_heading("Sources", level=1)
            for src in ctx.sources:
                p = doc.add_paragraph(style="List Number")
                p.add_run(f"{src['document_name']}\n").bold = True
                p.add_run(
                    f"Preview: {src['text_preview']}\n"
                    f"Relevance score: {src['score']:.3f}"
                )

        footer = doc.add_paragraph("DocVision OCR - AI-Powered Document Intelligence")
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.save(path)

    @staticmethod
    def _write_insights_docx(path: str, insights: Dict[str, str]):
        os.makedirs(os.path.dirname(path), exist_ok=True)
        doc = DocxDocument()
        title = doc.add_heading("DocVision OCR - Document Insights Report", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        section_titles = {
            "summary":          "Document Summary",
            "key_topics":       "Key Topics",
            "smart_notes":      "Smart Notes",
            "short_questions":  "Short Answer Questions",
            "long_questions":   "Long Answer Questions",
            "mcq":              "Multiple Choice Questions",
            "difficulty":       "Difficulty Analysis",
        }

        for key, title_text in section_titles.items():
            if key in insights and insights[key]:
                doc.add_heading(title_text, level=1)
                doc.add_paragraph(insights[key])

        doc.save(path)